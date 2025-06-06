import re
from docx import Document
#!!INSERISCI NOMI INTERLOCUTORI ALLE RIGHE 25 E 26
# francesca_pattern = re.compile(r'^FRANCESCA BELLESIA$', re.IGNORECASE) NOME INTERLOCUTORE PRINCIPALE
# alina_pattern = re.compile(r'^ALINA KOSTIUK$', re.IGNORECASE) NOME INTERLOCUTORE SECONDARIO
# nome= re.compile(r'^NOME COGNOME$', re.IGNORECASE)


def pulisci_docx_con_blocchi(input_file, output_file, soglia_caratteri=40):
    doc = Document(input_file)
    nuovo_doc = Document()

    # Prepara testo da tutti i paragrafi (riga per riga)
    righe = "\n".join(p.text for p in doc.paragraphs).splitlines()

    # Regex per orari
    pattern_orari = re.compile(
        r'\b\d{1,2}:\d{2}\b|'                # 0:07
        r'\b\d+\s*h\b|'                      # 1 h
        r'\b\d+\s*m\b|'                      # 2 m
        r'\b\d+\s*s\b|'                      # 30 s
        r'\b\d+\s*h\s*\d+\s*m\b|'            # 1 h 2 m
        r'\b\d+\s*m\s*\d+\s*s\b|'            # 2 m 30 s
        r'\b\d+\s*h\s*\d+\s*m\s*\d+\s*s\b',  # 1 h 2 m 30 s
        re.IGNORECASE
    )

    # Riconoscimento parlanti
    francesca_pattern = re.compile(r'^FRANCESCA BELLESIA$', re.IGNORECASE)
    alina_pattern = re.compile(r'^ALINA KOSTIUK$', re.IGNORECASE)

    # Frasi inutili
    #Se ne trovi altre aggiungile (semprecon doppio apice e virgola, non importa maiuscole o minuscole poiche IGNORECASE)
    frasi_inutili = {
        "ok", "hmm", "yes", "no", "yeah", "sure", "right", "i see", "thanks",
        "thank you", "interesting", "mh", "uh huh", "exactly", "mhm", "mm",
        "va bene", "perfetto"
    }

    # Stato
    speaker_corrente = None
    blocco_corrente = []

    def pulisci_riga(riga):
        return pattern_orari.sub('', riga).strip()

    def blocco_superfluo(righe):
        testo = " ".join(r.strip().lower() for r in righe)
        testo_pulito = re.sub(r'\W+', ' ', testo)
        parole = testo_pulito.split()
        tutte_inutili = all(w in frasi_inutili for w in parole)
        return len(testo_pulito) < soglia_caratteri and tutte_inutili

    def salva_blocco(speaker, righe_blocco):
        if speaker == "FRANCESCA BELLESIA":
            if blocco_superfluo(righe_blocco):
                print(f"❌ Blocco superfluo eliminato: {righe_blocco}")
                return
        nuovo_doc.add_paragraph(speaker)
        for r in righe_blocco:
            nuovo_doc.add_paragraph(r)
        nuovo_doc.add_paragraph()

    for riga in righe:
        riga_pulita = pulisci_riga(riga)
        if not riga_pulita:
            continue

        # Se trovi un parlante
        if francesca_pattern.match(riga_pulita):
            if speaker_corrente and blocco_corrente:
                salva_blocco(speaker_corrente, blocco_corrente)
            speaker_corrente = "FRANCESCA BELLESIA"
            blocco_corrente = []
        elif alina_pattern.match(riga_pulita):
            if speaker_corrente and blocco_corrente:
                salva_blocco(speaker_corrente, blocco_corrente)
            speaker_corrente = "ALINA KOSTIUK"
            blocco_corrente = []
        else:
            # Riga normale
            if speaker_corrente:
                blocco_corrente.append(riga_pulita)

    # Salva ultimo blocco
    if speaker_corrente and blocco_corrente:
        salva_blocco(speaker_corrente, blocco_corrente)

    nuovo_doc.save(output_file)
    print(f"✅ File pulito salvato: {output_file}")

#BLOCCO RIMOZIONE SPAZI E NOMI DUPLICATI

from docx import Document

def merge_and_deduplicate_speakers(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    current_speaker = None
    current_text = []
    last_written_speaker = None

    def flush_current():
        nonlocal last_written_speaker
        if current_speaker and current_text:
            if current_speaker != last_written_speaker:
                new_doc.add_paragraph(current_speaker)
                last_written_speaker = current_speaker
            new_doc.add_paragraph(" ".join(current_text).replace("\n", " ").strip())

    for para in doc.paragraphs:
        text = para.text.strip()

        # Se la riga è tutta maiuscola e breve, consideriamola un nome
        if text.isupper() and len(text.split()) <= 4:
            if text != current_speaker:
                flush_current()
                current_speaker = text
                current_text = []
        else:
            if text:
                current_text.append(text)

    # Flush finale
    flush_current()

    new_doc.save(output_path)
    print(f"File salvato: {output_path}")





# USO
if __name__ == "__main__":
    input_file = "ESEMPIO.docx"  # Il tuo file originale
    output_file = "ESEMPIO_FILTRATO.docx"
    pulisci_docx_con_blocchi(input_file, output_file, soglia_caratteri=40)
    merge_and_deduplicate_speakers(output_file, "res_finale.docx")

